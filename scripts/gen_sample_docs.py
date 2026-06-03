"""生成知识库测试用的 PDF / Word 样例文档。

PDF 依赖 reportlab（仅生成时需要，未列入项目依赖），运行方式：
    uv run --with reportlab python scripts/gen_sample_docs.py

输出目录：samples/
内容均为现有 Markdown 知识库未覆盖的运维主题，便于验证 PDF/Word 入库与检索。
"""

from pathlib import Path

from docx import Document as DocxDocument
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

OUTPUT_DIR = Path("samples")

# reportlab 内置 Adobe CJK 字体，无需外部字体文件即可渲染中文
CJK_FONT = "STSong-Light"


# ---- 文档内容（结构化为 标题 / 段落 / 表格，便于复用到 PDF 和 Word） ----

REDIS_DOC = {
    "title": "Redis 缓存雪崩与击穿处理手册",
    "blocks": [
        ("h", "告警说明"),
        ("p", "告警名称：RedisCacheMissSpike"),
        ("p", "触发条件：缓存命中率 5 分钟内从 95% 跌至 60% 以下，且数据库 QPS 同步飙升"),
        ("p", "严重等级：Critical　负责团队：缓存与中间件组"),
        ("h", "概念区分"),
        ("p", "缓存雪崩：大量 key 在同一时间集中过期，请求全部穿透到数据库。"),
        ("p", "缓存击穿：某个热点 key 过期瞬间，高并发请求同时打到数据库。"),
        ("p", "缓存穿透：查询根本不存在的 key，每次都绕过缓存访问数据库。"),
        ("h", "排查步骤"),
        ("p", "第一步：确认命中率下跌区间，定位是集中过期（雪崩）还是单 key（击穿）。"),
        ("p", "第二步：查看数据库连接数与慢查询，判断是否已被穿透流量压垮。"),
        ("p", "第三步：检查近期是否有批量写缓存操作使用了相同 TTL。"),
        ("h", "处置措施"),
        ("p", "雪崩：为 key 的过期时间叠加随机抖动（如 TTL + rand(0, 300s)），打散过期点。"),
        ("p", "击穿：对热点 key 加互斥锁或逻辑过期，只允许一个请求回源重建缓存。"),
        ("p", "穿透：对空结果缓存短 TTL 占位值，或引入布隆过滤器拦截非法 key。"),
        ("p", "兜底：开启限流与降级，保护数据库不被打挂。"),
        ("h", "升级路径"),
        ("table", [
            ["数据库 QPS 增幅", "操作"],
            ["< 2 倍", "观察，记录命中率曲线"],
            ["2–5 倍", "启用热点 key 互斥锁，通知业务负责人"],
            ["> 5 倍", "启动限流降级，拉起应急群"],
        ]),
    ],
}

KAFKA_DOC = {
    "title": "Kafka 消费积压处理手册",
    "blocks": [
        ("h", "告警说明"),
        ("p", "告警名称：KafkaConsumerLagHigh"),
        ("p", "触发条件：消费组 lag 持续 5 分钟超过 10 万条且持续增长"),
        ("p", "严重等级：Warning　负责团队：数据平台组"),
        ("h", "排查步骤"),
        ("p", "第一步：通过 kafka-consumer-groups.sh --describe 查看各分区 lag 分布。"),
        ("p", "第二步：判断是消费端变慢（处理耗时上升）还是生产端突增（写入流量激增）。"),
        ("p", "第三步：检查消费者实例是否有 rebalance 频繁、GC 停顿或下游依赖超时。"),
        ("h", "常见原因与处置"),
        ("p", "消费逻辑变慢：定位下游慢调用（数据库 / 第三方接口），优化或异步化处理。"),
        ("p", "分区分配不均：增加分区数或调整 key 哈希，使流量均衡到更多消费者。"),
        ("p", "消费者数量不足：在分区数允许范围内扩容消费者实例。"),
        ("p", "频繁 rebalance：排查心跳超时配置（session.timeout.ms / max.poll.interval.ms）。"),
        ("h", "临时缓解"),
        ("p", "对可丢弃的非关键消息，可临时调大 max.poll.records 提升批量吞吐。"),
        ("p", "若积压不可接受且消息可重放，考虑临时跳过到最新 offset 并记录丢失区间。"),
        ("h", "升级路径"),
        ("table", [
            ["lag 规模", "操作"],
            ["< 10 万", "观察，确认是否自然回落"],
            ["10 万–100 万", "扩容消费者，优化消费逻辑"],
            ["> 100 万", "启动应急，评估是否跳 offset 保时效"],
        ]),
    ],
}

SSL_DOC = {
    "title": "SSL/TLS 证书过期应急手册",
    "blocks": [
        ("h", "告警说明"),
        ("p", "告警名称：TLSCertExpiringSoon / TLSCertExpired"),
        ("p", "触发条件：证书剩余有效期 < 14 天（预警）或已过期（紧急）"),
        ("p", "严重等级：Warning / Critical　负责团队：基础设施运维组"),
        ("h", "影响范围"),
        ("p", "证书过期会导致 HTTPS 握手失败，客户端报 CERT_DATE_INVALID，全部依赖该域名的请求中断。"),
        ("p", "内部服务间 mTLS 证书过期会引发服务调用大面积失败，且不易被外部监控发现。"),
        ("h", "排查步骤"),
        ("p", "第一步：用 openssl s_client -connect host:443 查看证书有效期与签发链。"),
        ("p", "第二步：确认证书部署位置（负载均衡 / 网关 / 各服务实例），避免漏换。"),
        ("p", "第三步：检查证书是否由自动续期工具（如 cert-manager / certbot）管理及其失败原因。"),
        ("h", "处置措施"),
        ("p", "紧急续期：签发新证书并部署到所有终结 TLS 的节点，reload 而非重启以避免中断。"),
        ("p", "验证：重新执行 openssl 检查有效期，并从外部探测 HTTPS 握手恢复正常。"),
        ("p", "根因修复：补齐证书到期监控，确保自动续期任务有告警闭环。"),
        ("h", "升级路径"),
        ("table", [
            ["剩余有效期", "操作"],
            ["7–14 天", "提交续期工单，确认自动续期是否生效"],
            ["1–7 天", "人工介入续期，验证全部部署点"],
            ["已过期", "启动应急，最高优先级续期并恢复"],
        ]),
    ],
}

JVM_DOC = {
    "title": "JVM Full GC 频繁与 OOM 排查手册",
    "blocks": [
        ("h", "告警说明"),
        ("p", "告警名称：JVMFullGCFrequent / JVMHeapUsageHigh"),
        ("p", "触发条件：Full GC 频率 > 1 次/分钟，或老年代使用率持续 > 90%"),
        ("p", "严重等级：Critical　负责团队：应用支撑组"),
        ("h", "排查步骤"),
        ("p", "第一步：查看 GC 日志（-Xlog:gc*），确认 Full GC 频率、停顿时长与回收效果。"),
        ("p", "第二步：用 jstat -gcutil <PID> 1s 观察各代占用，判断是内存泄漏还是分配速率过高。"),
        ("p", "第三步：jmap -histo:live <PID> 查看存活对象 Top N，定位疑似泄漏的类。"),
        ("p", "第四步：必要时 jmap -dump 导出堆快照，用 MAT 分析支配树与 GC Roots 引用链。"),
        ("h", "常见原因"),
        ("p", "内存泄漏：静态集合无限增长、ThreadLocal 未清理、缓存无淘汰策略。"),
        ("p", "堆配置不当：-Xmx 偏小或新生代比例不合理，导致对象过早晋升老年代。"),
        ("p", "大对象 / 突发流量：一次性加载大结果集，瞬间撑爆堆内存。"),
        ("h", "处置措施"),
        ("p", "临时：重启实例恢复服务，并保留堆快照用于事后分析。"),
        ("p", "短期：调整堆与新生代参数，必要时扩容实例分摊压力。"),
        ("p", "根因：修复泄漏点，对大查询分页，对缓存设置容量上限与淘汰策略。"),
        ("h", "升级路径"),
        ("table", [
            ["老年代使用率", "操作"],
            ["80%–90%", "观察，准备堆快照"],
            ["> 90% 且频繁 FGC", "导出堆快照，评估重启"],
            ["已 OOM 崩溃", "启动应急，重启并定位泄漏根因"],
        ]),
    ],
}


def _build_pdf(doc_data: dict, output_path: Path) -> None:
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CJKTitle", parent=styles["Title"], fontName=CJK_FONT, fontSize=18, leading=24
    )
    heading_style = ParagraphStyle(
        "CJKHeading", parent=styles["Heading2"], fontName=CJK_FONT, fontSize=13, leading=18
    )
    body_style = ParagraphStyle(
        "CJKBody", parent=styles["BodyText"], fontName=CJK_FONT, fontSize=10.5, leading=16
    )

    flow = [Paragraph(doc_data["title"], title_style), Spacer(1, 6 * mm)]
    for kind, value in doc_data["blocks"]:
        if kind == "h":
            flow.append(Spacer(1, 3 * mm))
            flow.append(Paragraph(value, heading_style))
        elif kind == "p":
            flow.append(Paragraph(value, body_style))
        elif kind == "table":
            table = Table(value, hAlign="LEFT")
            table.setStyle(
                TableStyle(
                    [
                        ("FONTNAME", (0, 0), (-1, -1), CJK_FONT),
                        ("FONTSIZE", (0, 0), (-1, -1), 10),
                        ("GRID", (0, 0), (-1, -1), 0.5, (0.6, 0.6, 0.6)),
                        ("BACKGROUND", (0, 0), (-1, 0), (0.9, 0.9, 0.9)),
                        ("TOPPADDING", (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ]
                )
            )
            flow.append(Spacer(1, 2 * mm))
            flow.append(table)

    SimpleDocTemplate(str(output_path), pagesize=A4).build(flow)


def _build_docx(doc_data: dict, output_path: Path) -> None:
    document = DocxDocument()
    document.add_heading(doc_data["title"], level=0)
    for kind, value in doc_data["blocks"]:
        if kind == "h":
            document.add_heading(value, level=1)
        elif kind == "p":
            document.add_paragraph(value)
        elif kind == "table":
            rows, cols = len(value), len(value[0])
            table = document.add_table(rows=rows, cols=cols)
            table.style = "Light Grid Accent 1"
            for r, row in enumerate(value):
                for c, cell_text in enumerate(row):
                    table.rows[r].cells[c].text = cell_text
    document.save(str(output_path))


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    pdfmetrics.registerFont(UnicodeCIDFont(CJK_FONT))

    _build_pdf(REDIS_DOC, OUTPUT_DIR / "runbook_redis_cache_avalanche.pdf")
    _build_pdf(KAFKA_DOC, OUTPUT_DIR / "runbook_kafka_consumer_lag.pdf")
    _build_docx(SSL_DOC, OUTPUT_DIR / "runbook_ssl_cert_expiry.docx")
    _build_docx(JVM_DOC, OUTPUT_DIR / "runbook_jvm_full_gc.docx")

    for path in sorted(OUTPUT_DIR.iterdir()):
        print(f"generated: {path}  ({path.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
