from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

# 纯文本格式：直接按 UTF-8 读取
TEXT_EXTENSIONS = {".txt", ".md", ".markdown"}
# 二进制格式：抽取出的纯文本交给下游 split_text 处理
PDF_EXTENSION = ".pdf"
DOCX_EXTENSION = ".docx"

SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | {PDF_EXTENSION, DOCX_EXTENSION}


def extract_text(file_path: Path) -> str:
    """按扩展名把上传文件解析成纯文本。

    PDF/Word 抽取后即为无结构纯文本，下游 split_document 会走 split_text 路径。
    扫描件 PDF（无文本层）会抽出空字符串，由上层的空内容校验拦截。
    """
    suffix = file_path.suffix.lower()
    if suffix in TEXT_EXTENSIONS:
        # 非 UTF-8 会抛 UnicodeDecodeError，由上层转成对应的 400 提示
        return file_path.read_text(encoding="utf-8")
    if suffix == PDF_EXTENSION:
        return _extract_pdf(file_path)
    if suffix == DOCX_EXTENSION:
        return _extract_docx(file_path)
    raise ValueError(f"不支持的文件类型：{suffix}")


def _extract_pdf(file_path: Path) -> str:
    reader = PdfReader(str(file_path))
    # 逐页抽取文本层；某页无文本时 extract_text 返回 None，用空串兜底
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(file_path: Path) -> str:
    document = DocxDocument(str(file_path))
    # 段落正文 + 表格单元格文本，避免运维文档中的表格内容丢失
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)
